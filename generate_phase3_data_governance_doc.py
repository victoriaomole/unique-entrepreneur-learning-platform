import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------- CONFIG ----------
ORG_NAME = "UNIQUE ENTREPRENEUR SOLUTIONS LIMITED"
FOOTER_TEXT = (
    "Omole Victoria Oluwatosin — Business Analyst | Data Analyst | "
    "Data Management Professional | Frontend & Backend Engineer"
)
BLUE_GREY = RGBColor(30, 55, 90)
LIGHT_GREY = "F0F0F0"

BASE_DIR = r"C:\Users\victo\kickoff_project\02_Phase2_Requirements_and_Governance"
os.makedirs(BASE_DIR, exist_ok=True)

OUTPUT_PATH = os.path.join(
    BASE_DIR,
    "UniqueEntrepreneur_Phase3_Data_Governance_Implementation.docx"
)

# ---------- HELPERS ----------

def shade_cell(cell, fill_hex=LIGHT_GREY):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

def set_header_footer(section):
    # Header
    header = section.header
    if not header.paragraphs:
        header.add_paragraph()
    hp = header.paragraphs[0]
    hp.text = ""
    hr = hp.add_run(ORG_NAME)
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr.font.size = Pt(11)
    hr.font.bold = True
    hr.font.color.rgb = BLUE_GREY

    # Footer
    footer = section.footer
    if not footer.paragraphs:
        footer.add_paragraph()
    fp = footer.paragraphs[0]
    fp.text = ""
    fr = fp.add_run(FOOTER_TEXT)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr.font.size = Pt(9)
    fr.font.color.rgb = BLUE_GREY

def add_title(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = BLUE_GREY

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = BLUE_GREY

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = BLUE_GREY
    doc.add_paragraph()

def add_text_area(doc, label):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.font.bold = True
    r.font.color.rgb = BLUE_GREY

    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    p2 = cell.paragraphs[0]
    r2 = p2.add_run("Click to type")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(120, 120, 120)
    shade_cell(cell)
    doc.add_paragraph()

def add_matrix_table(doc, title, headers, rows):
    add_section_heading(doc, title)
    table = doc.add_table(rows=1, cols=len(headers))

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BLUE_GREY

    # Body rows
    for row in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            para = row_cells[i].paragraphs[0]
            run = para.add_run(value)
            run.font.size = Pt(10)
            if i == 0:
                run.font.bold = True
                run.font.color.rgb = BLUE_GREY
            else:
                run.font.color.rgb = RGBColor(120, 120, 120)
            shade_cell(row_cells[i])

    doc.add_paragraph()

# ---------- DOCUMENT CREATION ----------

doc = Document()

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
style.font.size = Pt(10)

set_header_footer(doc.sections[0])

# Cover
add_title(doc, "Phase 3B — Data Governance Implementation Guide")
add_subtitle(doc, "Linking Policy (Phase 2) to Technical Design (Phase 3)")
doc.add_paragraph()

# Metadata
meta = doc.add_table(rows=3, cols=2)
labels = ["Project Name:", "Organisation:", "Document Type:"]
values = ["Unique Entrepreneur Literacy Hub", ORG_NAME, "Data Governance Implementation (Technical Mapping)"]
for i in range(3):
    c1, c2 = meta.rows[i].cells
    r1 = c1.paragraphs[0].add_run(labels[i])
    r1.font.bold = True
    r1.font.color.rgb = BLUE_GREY
    r2 = c2.paragraphs[0].add_run(values[i])
    r2.font.color.rgb = RGBColor(120,120,120)
    shade_cell(c2)
doc.add_paragraph()

# 1. Purpose
add_section_heading(doc, "1. Purpose")
add_text_area(doc,
    "Explain how this document translates the high-level Data Governance Framework "
    "into concrete RBAC, RLS, logging, and retention configurations in the platform."
)

# 2. RBAC Implementation Matrix
add_matrix_table(
    doc,
    "2. RBAC Implementation Matrix",
    ["Role", "System Capabilities (Examples)", "Technical Notes"],
    [
        ["Platform Admin",
         "Manage tenants, global settings, view cross-tenant metrics (with safeguards).",
         "Django group/claim: platform_admin = True. Access via admin-only endpoints."],
        ["Org Admin",
         "Manage schools, classes, users within their org; view org-level reports.",
         "Filter by org_id from JWT; DRF permissions restrict to that org_id."],
        ["School Admin",
         "Manage classes, enrolments, view attendance for their school.",
         "Filter by school_id; cannot change org-level settings."],
        ["Instructor",
         "Create/manage their courses; see learners only on those courses.",
         "Ownership via instructor_id on Course; queries scoped accordingly."],
        ["Student",
         "Access enrolled courses, own progress and certificates only.",
         "All queries filtered by user_id; no cross-user access."],
        ["DMP / Compliance",
         "Read-only access to logs, retention configs, exports (where authorised).",
         "Dedicated permission; access only via secure audit/report endpoints."],
    ]
)

# 3. RLS Strategy per Table
add_matrix_table(
    doc,
    "3. Row-Level Security (RLS) Strategy",
    ["Table", "RLS Rule Summary", "Notes"],
    [
        ["organisation",
         "No RLS for platform admins; restricted views for others.",
         "Normally only platform_admin sees all orgs."],
        ["school",
         "school.org_id must match current_org_id.",
         "Org Admin & above only."],
        ["course",
         "course.org_id is NULL (public) or = current_org_id.",
         "Instructors see owned courses; org admins see all in org."],
        ["enrollment",
         "enrollment.org_id = current_org_id AND (role-based constraints).",
         "Students see their own; org/school admins see within scope."],
        ["audit_log",
         "Restricted to DMP/Platform Admin; may be per-org.",
         "No general user access."],
    ]
)

# 4. Data Flows & Lineage (Template)
add_section_heading(doc, "4. Data Flows & Lineage")
add_text_area(doc,
    "Describe key flows: registration, enrolment, learning events, quiz submissions, payments, "
    "and how data moves between frontend, API, DB, storage, and external services."
)

# 5. Logging & Audit Requirements (Technical)
add_section_heading(doc, "5. Logging & Audit Implementation")
add_text_area(doc,
    "List which events are logged (logins, SSO, role changes, config changes, payouts, grade changes), "
    "where logs are stored, and how they are accessed securely by DMP/Platform Admin."
)

# 6. Retention & Deletion Implementation
add_section_heading(doc, "6. Retention & Deletion Implementation")
add_matrix_table(
    doc,
    "6.1 Automation Rules Template",
    ["Data Category", "Retention Logic (System)", "Deletion / Anonymisation Approach"],
    [
        ["User Profile",
         "Keep active; after X years of inactivity, anonymise or delete.",
         "Background job checks last_login, flags for anonymisation."],
        ["Learning Records",
         "Retain for X years post-completion.",
         "Soft delete vs. full delete configurable per org."],
        ["Payments",
         "Minimum 7 years.",
         "Hard delete disabled; only restricted access."],
        ["Audit Logs",
         "X years.",
         "Stored append-only; purged on schedule."],
    ]
)

# 7. Data Quality Controls
add_section_heading(doc, "7. Data Quality Controls")
add_text_area(doc,
    "Define validation rules (required fields, allowed values), duplicate detection, and monitoring checks. "
    "Describe how issues are surfaced to Org Admins or support."
)

# 8. Access to Exports & Subject Rights
add_section_heading(doc, "8. Data Exports & Subject Rights")
add_text_area(doc,
    "Describe how a DPO/DMP or Org Admin can trigger exports for a user (subject access), "
    "execute erasure requests, and how these actions are audited."
)

# 9. Environment & Keys Governance
add_section_heading(doc, "9. Secrets & Environment Governance")
add_text_area(doc,
    "Document rules for managing API keys, JWT secrets, DB credentials, S3 keys, and admin accounts across "
    "Dev / Test / Prod. Include who can access which environment."
)

# 10. Approval & Review
add_section_heading(doc, "10. Approval & Review")
add_text_area(doc,
    "Capture approvals (Platform Admin, DMP, Legal) and define how often this implementation guide "
    "is reviewed and updated.")
    
# Save
doc.save(OUTPUT_PATH)
print(f"✅ Phase 3B Data Governance Implementation template created at:\n{OUTPUT_PATH}")
