import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------- CONFIG ----------
ORG_NAME = "UNIQUE ENTREPRENEUR SOLUTIONS LIMITED"
FOOTER_TEXT = (
    "Omole Victoria Oluwatosin — Business Analyst | Data Analyst | "
    "Data Management Professional | Frontend & Backend Engineer"
)
BLUE_GREY = RGBColor(30, 55, 90)
LIGHT_GREY = "F0F0F0"

BASE_DIR = r"C:\Users\victo\kickoff_project\02_Phase2_Requirements_and_Governance"
os.makedirs(BASE_DIR, exist_ok=True)

OUTPUT_PATH = os.path.join(
    BASE_DIR,
    "UniqueEntrepreneur_phase3_Architecture_and_ERD.docx"
)

# ---------- HELPERS ----------

def shade_cell(cell, fill_hex=LIGHT_GREY):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

def set_header_footer(section):
    header = section.header
    if not header.paragraphs:
        header.add_paragraph()
    hp = header.paragraphs[0]
    hp.text = ""
    hr = hp.add_run(ORG_NAME)
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr.font.size = Pt(11)
    hr.font.bold = True
    hr.font.color.rgb = BLUE_GREY

    footer = section.footer
    if not footer.paragraphs:
        footer.add_paragraph()
    fp = footer.paragraphs[0]
    fp.text = ""
    fr = fp.add_run(FOOTER_TEXT)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr.font.size = Pt(9)
    fr.font.color.rgb = BLUE_GREY

def add_title(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = BLUE_GREY

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = BLUE_GREY

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = BLUE_GREY
    doc.add_paragraph()

def add_text_area(doc, label):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.font.bold = True
    r.font.color.rgb = BLUE_GREY
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    p2 = cell.paragraphs[0]
    r2 = p2.add_run("Click to type")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(120, 120, 120)
    shade_cell(cell)
    doc.add_paragraph()

def add_kv_table(doc, title, rows):
    add_section_heading(doc, title)
    table = doc.add_table(rows=0, cols=2)
    for label, value in rows:
        row = table.add_row().cells
        lp = row[0].paragraphs[0].add_run(label)
        lp.font.bold = True
        lp.font.color.rgb = BLUE_GREY
        vp = row[1].paragraphs[0].add_run(value)
        vp.font.color.rgb = RGBColor(120,120,120)
        shade_cell(row[1])
    doc.add_paragraph()

# ---------- DOCUMENT CREATION ----------

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
style.font.size = Pt(10)

set_header_footer(doc.sections[0])

add_title(doc, "System Architecture & ERD")
add_subtitle(doc, "Unique Entrepreneur Literacy Hub — Phase 3 Design Baseline")
doc.add_paragraph()

# 1. Overview
add_section_heading(doc, "1. Overview")
add_text_area(doc, "Briefly describe the platform purpose, tenants (co-ops, SMEs, schools), "
                   "and target regions (UK & Nigeria).")

# 2. Architecture Summary
add_section_heading(doc, "2. Architecture Summary")
add_kv_table(doc, "2.1 Technology Stack", [
    ("Frontend", "Next.js (App Router), React, TypeScript, Tailwind, shadcn/ui"),
    ("Backend", "Django, Django REST Framework, JWT auth"),
    ("Database", "PostgreSQL with Row-Level Security"),
    ("Cache & Queue", "Redis + Celery"),
    ("Storage/CDN", "AWS S3 + CloudFront"),
    ("Video", "S3 + transcoding or Vimeo/Wistia"),
    ("Infra", "Docker, CI/CD (GitHub Actions), AWS/Render/Vercel/Fly.io"),
])

add_text_area(doc, "2.2 High-Level Component Diagram (describe services: Web, API, DB, Cache, Storage, Integrations).")

# 3. Multi-Tenancy & Isolation
add_section_heading(doc, "3. Multi-Tenancy & Data Isolation")
add_text_area(doc, "Explain org → school → class hierarchy, tenant resolution "
                   "(domain/subdomain/claims), and isolation rules by org_id / school_id.")

# 4. ERD (Entities & Relations)
add_section_heading(doc, "4. Entity-Relationship Model (ERD)")
add_text_area(doc, "List core entities (Organisation, School, Class, User, RoleAssignment, Course, Section, Lesson, "
                   "Enrollment, Quiz, Submission, Certificate, Order, Payment, Payout, Coupon, AuditLog). "
                   "Attach diagram or maintain link to draw.io/Lucidchart.")

# 5. Row-Level Security (RLS)
add_section_heading(doc, "5. Row-Level Security Strategy")
add_text_area(doc, "Describe how PostgreSQL RLS is applied (e.g. policies on course, enrollment, etc.) "
                   "based on current_org_id, roles, and tenant context from JWT.")

# 6. API Design (Alignment with OpenAPI)
add_section_heading(doc, "6. API Design Overview")
add_text_area(doc, "Summarise main endpoints (auth, orgs, schools, courses, enrolments, quizzes, reports), "
                   "use of pagination, filtering, and versioning (/api/v1).")

# 7. Non-Functional Requirements Mapping
add_section_heading(doc, "7. Non-Functional Requirements Mapping")
add_text_area(doc, "Explain how the chosen architecture meets performance, scalability, observability, "
                   "security, and availability targets.")

# 8. Integration Points
add_section_heading(doc, "8. Integrations")
add_text_area(doc, "List integrations: SSO (OAuth/SAML), payment gateways, email provider, analytics, video hosting.")

# 9. Security, Privacy & Data Governance Link
add_section_heading(doc, "9. Security & Data Governance Alignment")
add_text_area(doc, "Reference Data Governance Framework; show how RBAC, RLS, encryption, audit logs "
                   "and retention are enforced technically.")

# 10. Open Questions & Decisions Log
add_section_heading(doc, "10. Open Questions & Design Decisions")
add_text_area(doc, "Track pending decisions (e.g. final video provider, final hosting choice, etc.).")

doc.save(OUTPUT_PATH)
print(f"✅ Architecture & ERD design template created at:\n{OUTPUT_PATH}")
