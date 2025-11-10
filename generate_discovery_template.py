from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------- CONFIG ----------
ORG_NAME = "UNIQUE ENTREPRENEUR SOLUTIONS LIMITED"
FOOTER_TEXT = ("Omole Victoria Oluwatosin — Business Analyst | Data Analyst | "
               "Data Management Professional | Frontend & Backend Engineer")

BLUE_GREY = RGBColor(30, 55, 90)


# ---------- HELPERS ----------

def shade_cell(cell, fill_hex="F0F0F0"):
    """Apply light background shading to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def set_header_footer(section):
    # Header
    header = section.header
    if not header.paragraphs:
        header.add_paragraph()
    header_para = header.paragraphs[0]
    header_para.text = ORG_NAME
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header_para.runs:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = BLUE_GREY

    # Footer
    footer = section.footer
    if not footer.paragraphs:
        footer.add_paragraph()
    footer_para = footer.paragraphs[0]
    footer_para.text = FOOTER_TEXT
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = BLUE_GREY


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = BLUE_GREY


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE_GREY


def add_label_input_row(table, label_text, placeholder="Click to type"):
    row = table.add_row()
    label_cell = row.cells[0]
    input_cell = row.cells[1]

    lp = label_cell.paragraphs[0]
    lr = lp.add_run(label_text)
    lr.font.bold = True
    lr.font.size = Pt(10)
    lr.font.color.rgb = BLUE_GREY

    ip = input_cell.paragraphs[0]
    ir = ip.add_run(placeholder)
    ir.font.size = Pt(10)
    ir.font.color.rgb = RGBColor(120, 120, 120)
    shade_cell(input_cell)


def add_section_break(doc):
    doc.add_page_break()


def create_generic_header_block(doc, title_text):
    add_title(doc, title_text)
    add_subtitle(doc, "Discovery Interview Template – Phase 1: Discovery & Scoping")
    doc.add_paragraph()  # spacer

    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Interview Details"
    hdr_cells[1].text = ""

    hp = hdr_cells[0].paragraphs[0]
    hr = hp.add_run("Interview Details")
    hr.font.bold = True
    hr.font.size = Pt(10)
    hr.font.color.rgb = BLUE_GREY

    add_label_input_row(table, "Interview ID:")
    add_label_input_row(table, "Date / Time:")
    add_label_input_row(table, "Interviewee Name:")
    add_label_input_row(table, "Organisation / School:")
    add_label_input_row(table, "Role / Title:")
    add_label_input_row(table, "Conducted by:", "Omole Victoria Oluwatosin")
    add_label_input_row(table, "Recording Consent Given (Yes / No):")
    add_label_input_row(table, "Summary Tags (e.g., Payments, UX, Governance):")

    doc.add_paragraph()  # spacer


def add_question_block(doc, heading, questions):
    p = doc.add_paragraph()
    run = p.add_run(heading)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE_GREY

    for q in questions:
        # question
        qp = doc.add_paragraph()
        qp.paragraph_format.left_indent = Pt(6)
        qr = qp.add_run(f"- {q}")
        qr.font.size = Pt(10)

        # answer area
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        ap = cell.paragraphs[0]
        ar = ap.add_run("Click here to type response")
        ar.font.size = Pt(10)
        ar.font.color.rgb = RGBColor(120, 120, 120)
        shade_cell(cell)

    doc.add_paragraph()  # spacer


def add_summary_block(doc):
    p = doc.add_paragraph()
    r = p.add_run("Summary & Key Insights")
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = BLUE_GREY

    themes = [
        "User Experience",
        "Functionality",
        "Data & Analytics",
        "Governance & Security",
        "Other Notes",
    ]
    headers = ["Theme", "Key Insights / Quotes", "Opportunities / Actions"]

    rows = len(themes) + 1
    table = doc.add_table(rows=rows, cols=3)

    # header row
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        hp = hdr_cells[i].paragraphs[0]
        hr = hp.add_run(text)
        hr.font.bold = True
        hr.font.size = Pt(10)
        hr.font.color.rgb = BLUE_GREY

    # theme rows
    for i, theme in enumerate(themes, start=1):
        row_cells = table.rows[i].cells

        tp = row_cells[0].paragraphs[0]
        tr = tp.add_run(theme)
        tr.font.bold = True
        tr.font.size = Pt(10)
        tr.font.color.rgb = BLUE_GREY

        for j in (1, 2):
            cp = row_cells[j].paragraphs[0]
            cr = cp.add_run("Click to type")
            cr.font.size = Pt(10)
            cr.font.color.rgb = RGBColor(120, 120, 120)
            shade_cell(row_cells[j])


# ---------- DOCUMENT CREATION ----------

doc = Document()

# default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
style.font.size = Pt(10)

set_header_footer(doc.sections[0])

# A – Org Admin
create_generic_header_block(doc, "Section A – Organisation Admin / Owner")
add_question_block(doc, "Key Questions", [
    "How is your organisation currently delivering and tracking online learning?",
    "What challenges do you face managing multiple schools or departments?",
    "What permissions or controls are most important for your admins?",
    "How do you handle enrolments, attendance, and reporting today?",
    "Which integrations (e.g., Microsoft Entra ID, Google Workspace) do you rely on?",
    "How do you ensure compliance and data privacy currently?",
    "Which analytics or reports are essential for you?",
    "What would make you adopt this new platform?",
])
add_summary_block(doc)

add_section_break(doc)
set_header_footer(doc.sections[-1])

# B – Instructor
create_generic_header_block(doc, "Section B – Instructor / Teacher")
add_question_block(doc, "Key Questions", [
    "Describe your typical process for planning and creating a course.",
    "Which content formats do you use (video, PDF, quizzes, assignments, live sessions)?",
    "How do you currently manage assessments and feedback?",
    "What features would simplify your teaching workflow?",
    "How would you like to track student progress and performance?",
    "What monetisation or pricing options do you need?",
    "How should certificates and feedback be managed?",
    "What frustrates you about existing platforms?",
])
add_summary_block(doc)

add_section_break(doc)
set_header_footer(doc.sections[-1])

# C – Student
create_generic_header_block(doc, "Section C – Student / Learner")
add_question_block(doc, "Key Questions", [
    "What motivates you to enrol on an online course?",
    "Which learning formats do you prefer (video, reading, interactive, live)?",
    "Describe a recent good or poor experience using an online learning platform.",
    "What frustrates you about enrolment or navigation?",
    "Do you need mobile or offline access?",
    "How should your progress and certificates be displayed or shared?",
    "Which types of notifications are useful (deadlines, new content, reminders)?",
    "What does a successful learning experience look like for you?",
])
add_summary_block(doc)

add_section_break(doc)
set_header_footer(doc.sections[-1])

# D – Data Management / Compliance
create_generic_header_block(doc, "Section D – Data Management / Compliance")
add_question_block(doc, "Key Questions", [
    "Which data protection and privacy regulations apply to your organisation?",
    "How long should learner and course data be retained?",
    "Who should own data within the platform for your organisation?",
    "What are your requirements for consent, audit logs, and data export?",
    "How do you currently manage subject access and deletion requests?",
    "Which roles should have access to what categories of data?",
    "What backup and disaster recovery expectations do you have?",
    "What controls or reports are needed for audits?",
])
add_summary_block(doc)

add_section_break(doc)
set_header_footer(doc.sections[-1])

# E – Executive Sponsor
create_generic_header_block(doc, "Section E – Executive Sponsor / Product Owner")
add_question_block(doc, "Key Questions", [
    "What business outcomes define success for this platform?",
    "Which problems are we solving that are most urgent?",
    "What are the top three KPIs you expect to see post-launch?",
    "What does a realistic, valuable MVP include?",
    "What are the critical timeline or budget boundaries?",
    "Are there strategic integrations or partners we must support?",
    "What risks or failure modes concern you most?",
])
add_summary_block(doc)

add_section_break(doc)
set_header_footer(doc.sections[-1])

# F – Technical Lead
create_generic_header_block(doc, "Section F – Technical / Engineering Lead")
add_question_block(doc, "Key Questions", [
    "What preferred technologies or platforms should we align with?",
    "Are there existing systems or APIs we need to integrate with (SSO, LMS, SIS, payments)?",
    "What security or hosting constraints must be followed?",
    "How should environments (dev/test/prod) be structured?",
    "Any specific logging, monitoring, or observability requirements?",
    "What are the main technical risks you foresee?",
])
add_summary_block(doc)

# Save
filename = "UniqueEntrepreneur_Discovery_Interview_Template.docx"
doc.save(filename)
print(f"✅ Created: {filename}")

