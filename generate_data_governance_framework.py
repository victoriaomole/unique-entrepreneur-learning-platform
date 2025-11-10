import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------- CONFIG ----------
ORG_NAME = "UNIQUE ENTREPRENEUR SOLUTIONS LIMITED"
FOOTER_TEXT = (
    "Omole Victoria Oluwatosin — Business Analyst | Data Analyst | "
    "Data Management Professional | Frontend & Backend Engineer"
)
BLUE_GREY = RGBColor(30, 55, 90)
LIGHT_GREY = "F0F0F0"

BASE_DIR = r"C:\Users\victo\kickoff_project\02_Phase2_Requirements_and_Governance"
os.makedirs(BASE_DIR, exist_ok=True)

OUTPUT_PATH = os.path.join(
    BASE_DIR,
    "UniqueEntrepreneur_Data_Governance_Framework.docx"
)

# ---------- HELPERS ----------

def shade_cell(cell, fill_hex=LIGHT_GREY):
    """Apply light background shading to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

def set_header_footer(section):
    # Header
    header = section.header
    if not header.paragraphs:
        header.add_paragraph()
    hp = header.paragraphs[0]
    hp.text = ""
    hr = hp.add_run(ORG_NAME)
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr.font.size = Pt(11)
    hr.font.bold = True
    hr.font.color.rgb = BLUE_GREY

    # Footer
    footer = section.footer
    if not footer.paragraphs:
        footer.add_paragraph()
    fp = footer.paragraphs[0]
    fp.text = ""
    fr = fp.add_run(FOOTER_TEXT)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr.font.size = Pt(9)
    fr.font.color.rgb = BLUE_GREY

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = BLUE_GREY

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = BLUE_GREY

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = BLUE_GREY
    doc.add_paragraph()

def add_text_area(doc, label):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.font.bold = True
    r.font.color.rgb = BLUE_GREY

    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    p2 = cell.paragraphs[0]
    r2 = p2.add_run("Click to type")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(120, 120, 120)
    shade_cell(cell)

    doc.add_paragraph()

def add_matrix_table(doc, title, headers, rows_hints):
    add_section_heading(doc, title)
    table = doc.add_table(rows=1, cols=len(headers))
    # header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        pr = cell.paragraphs[0].add_run(h)
        pr.font.bold = True
        pr.font.size = Pt(10)
        pr.font.color.rgb = BLUE_GREY
    # body rows (with hints, editable)
    for row_hint in rows_hints:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_hint):
            para = row_cells[i].paragraphs[0]
            run = para.add_run(text)
            run.font.size = Pt(10)
            if i == 0:
                run.font.bold = True
                run.font.color.rgb = BLUE_GREY
            else:
                run.font.color.rgb = RGBColor(120, 120, 120)
            shade_cell(row_cells[i])
    doc.add_paragraph()

# ---------- DOCUMENT CREATION ----------

doc = Document()

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
style.font.size = Pt(10)

set_header_footer(doc.sections[0])

# Cover
add_title(doc, "Data Governance Framework")
add_subtitle(doc, "Phase 2 — Requirements & Governance Alignment")
doc.add_paragraph()

# Metadata
meta_table = doc.add_table(rows=3, cols=2)
labels = ["Project Name:", "Organisation:", "Version:"]
values = ["Unique Entrepreneur Literacy Hub", ORG_NAME, "1.0 (Template)"]
for i in range(3):
    c1, c2 = meta_table.rows[i].cells
    r1 = c1.paragraphs[0].add_run(labels[i])
    r1.font.bold = True
    r1.font.color.rgb = BLUE_GREY
    r2 = c2.paragraphs[0].add_run(values[i])
    r2.font.color.rgb = RGBColor(120,120,120)
    shade_cell(c2)
doc.add_paragraph()

# 1. Purpose & Scope
add_section_heading(doc, "1. Purpose & Scope")
add_text_area(doc, "Define the objectives of data governance for the platform, "
                   "covering multi-tenant co-ops, agribusiness, recycling SMEs in the UK & Nigeria.")

# 2. Data Governance Principles
add_section_heading(doc, "2. Data Governance Principles")
add_text_area(doc, "Document core principles, e.g. accountability, transparency, security, privacy, data minimisation, "
                   "and responsible use of learner and financial data.")

# 3. Roles & Responsibilities (RACI-style overview)
add_matrix_table(
    doc,
    "3. Roles & Responsibilities",
    ["Role", "Key Responsibilities", "Notes"],
    [
        ["Platform Admin",
         "Global configuration, tenant provisioning, enforcement of policies.",
         "Click to refine responsibilities."],
        ["Org Admin / Co-operative Lead",
         "Manage users, courses, and local compliance within their organisation.",
         "Click to specify what they own vs. platform."],
        ["Data Management Professional (DMP)",
         "Define policies, oversee data quality, approve retention & exports.",
         "Click to list named individuals or teams."],
        ["Technical Lead",
         "Implements security controls, backups, integrations.",
         "Click to reference architecture documents."],
    ]
)

# 4. Data Inventory & Classification
add_section_heading(doc, "4. Data Inventory & Classification")
add_matrix_table(
    doc,
    "4.1 Key Data Entities",
    ["Data Entity", "Description", "Classification", "Owner"],
    [
        ["User / Member Profile",
         "Names, contact info, role, org/school mapping.",
         "Confidential (Personal Data)",
         "Org Admin / Platform Admin"],
        ["Course & Content",
         "Titles, modules, media, metadata.",
         "Internal / Public (varies by visibility)",
         "Instructor / Org Admin"],
        ["Learning Records",
         "Enrolments, quiz scores, completions, certificates.",
         "Confidential (Educational Data)",
         "Platform + Org Admin"],
        ["Payments & Payouts",
         "Transactions, payouts, invoices.",
         "Highly Confidential (Financial Data)",
         "Platform Admin / Finance"],
    ]
)

# 5. Data Access & Row-Level Security (RLS)
add_section_heading(doc, "5. Data Access & Row-Level Security (RLS)")
add_text_area(doc,
              "Describe how access is restricted:\n"
              "- Isolation by org_id and (if applicable) school_id.\n"
              "- Platform Admin: cross-tenant read (with controls).\n"
              "- Org Admin: access limited to their organisation.\n"
              "- Instructors: access only to their courses and enrolled learners.\n"
              "- Students: access only to their own records.\n"
              "- DMP: governed access for audits and compliance.")

# 6. Data Retention & Deletion
add_section_heading(doc, "6. Data Retention & Deletion")
add_matrix_table(
    doc,
    "6.1 Retention Rules (Template)",
    ["Data Category", "Retention Period", "Notes / Legal Basis"],
    [
        ["User Accounts",
         "Active + X years after last activity",
         "Configure per org; align with GDPR/NDPR."],
        ["Learning Records",
         "X years after completion",
         "Sufficient for reporting and audits."],
        ["Payments & Invoices",
         "7+ years",
         "Financial regulations."],
        ["Audit Logs",
         "X years",
         "Support investigations and compliance."],
    ]
)

# 7. Data Quality Management
add_section_heading(doc, "7. Data Quality Management")
add_text_area(doc,
              "Define validation rules, duplicate handling, mandatory fields, "
              "and responsibilities for correcting inaccurate records.")

# 8. Security & Privacy Controls
add_section_heading(doc, "8. Security & Privacy Controls")
add_text_area(doc,
              "Document controls: HTTPS, encryption at rest, RBAC, RLS, backups, "
              "malware scanning for uploads, incident response, DPIA requirements, "
              "and how GDPR/NDPR rights (access, rectification, erasure) are supported.")

# 9. Audit, Monitoring & Reporting
add_section_heading(doc, "9. Audit, Monitoring & Reporting")
add_text_area(doc,
              "Define which events are logged (logins, role changes, data exports, grade changes), "
              "how often logs are reviewed, and who has access.")

# 10. Data Lineage & Integrations
add_section_heading(doc, "10. Data Lineage & Integrations")
add_text_area(doc,
              "List integrated systems (SSO, payments, video hosting, analytics) and describe "
              "how data flows between them, including ownership and responsibilities.")

# 11. Approval & Review
add_section_heading(doc, "11. Approval & Review")
add_text_area(doc,
              "Capture approvals (DMP, Platform Admin, Legal) and define review frequency "
              "for this framework (e.g. annually or after major platform changes).")

# Save
doc.save(OUTPUT_PATH)
print(f"✅ Data Governance Framework template created at:\n{OUTPUT_PATH}")

