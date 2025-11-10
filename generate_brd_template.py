from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

# ---------- CONFIG ----------
ORG_NAME = "UNIQUE ENTREPRENEUR SOLUTIONS LIMITED"
FOOTER_TEXT = (
    "Omole Victoria Oluwatosin — Business Analyst | Data Analyst | "
    "Data Management Professional | Frontend & Backend Engineer"
)
BLUE_GREY = RGBColor(30, 55, 90)
LIGHT_GREY_HEX = "F0F0F0"

BASE_DIR = r"C:\Users\victo\kickoff_project"
OUTPUT_PATH = os.path.join(
    BASE_DIR,
    "01_Phase1_Discovery_and_Scoping",
    "UniqueEntrepreneur_BRD_Template.docx",
)

# ---------- HELPERS ----------

def shade_cell(cell, fill_hex=LIGHT_GREY_HEX):
    """Apply light background shading to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

def set_header_footer(section):
    # Header
    header = section.header
    if not header.paragraphs:
        header.add_paragraph()
    hp = header.paragraphs[0]
    hp.text = ""
    run = hp.add_run(ORG_NAME)
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = BLUE_GREY

    # Footer
    footer = section.footer
    if not footer.paragraphs:
        footer.add_paragraph()
    fp = footer.paragraphs[0]
    fp.text = ""
    frun = fp.add_run(FOOTER_TEXT)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun.font.size = Pt(9)
    frun.font.color.rgb = BLUE_GREY

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = BLUE_GREY

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = BLUE_GREY

def add_label_input(doc, label, placeholder="Click to type"):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True

    c1, c2 = table.rows[0].cells
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(label)
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = BLUE_GREY

    p2 = c2.paragraphs[0]
    r2 = p2.add_run(placeholder)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(120, 120, 120)
    shade_cell(c2)

    doc.add_paragraph()  # spacer

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = BLUE_GREY
    doc.add_paragraph()

def add_text_area(doc, label):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.font.bold = True
    r.font.color.rgb = BLUE_GREY

    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    p2 = cell.paragraphs[0]
    r2 = p2.add_run("Click to type")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(120, 120, 120)
    shade_cell(cell)

    doc.add_paragraph()  # spacer

# ---------- DOCUMENT CREATION ----------

doc = Document()

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
style.font.size = Pt(10)

# Header & footer for first (only) section
set_header_footer(doc.sections[0])

# Title block
add_title(doc, "Business Requirements Document (BRD)")
add_subtitle(doc, "Phase 2 — Requirements & Governance Alignment")
doc.add_paragraph()

# Metadata
add_label_input(doc, "Project Name:", "Unique Entrepreneur Literacy Hub")
add_label_input(doc, "Organisation:", ORG_NAME)
add_label_input(doc, "Prepared By:", "Omole Victoria Oluwatosin")
add_label_input(doc, "Date:", "Click to type")
add_label_input(doc, "Version:", "1.0")

# Sections
add_section_heading(doc, "1. Purpose & Background")
add_text_area(doc, "Explain the business context, target users (cooperatives, small agribusinesses, recycling SMEs), and why this platform is needed.")

add_section_heading(doc, "2. Scope")
add_text_area(doc, "Define in-scope modules (multi-tenant LMS, course delivery, governance) and out-of-scope items (future AI, mobile apps, etc.).")

add_section_heading(doc, "3. Stakeholders & Roles")
add_text_area(doc, "List Platform Admin, Org Admin, School Admin, Instructor, Learner, DMP, Technical Lead, Support/Moderator, etc., with responsibilities.")

add_section_heading(doc, "4. Functional Requirements")
add_text_area(doc, "Capture FR-01, FR-02, etc. For each: feature, description, priority, dependencies.")

add_section_heading(doc, "5. Non-Functional Requirements")
add_text_area(doc, "Document performance targets, security (JWT, RLS), availability, localisation, accessibility, scalability, and tech constraints.")

add_section_heading(doc, "6. Data Governance & Compliance")
add_text_area(doc, "Define data ownership, GDPR/NDPR alignment, RLS rules, consent, audit logs, retention periods, and breach handling.")

add_section_heading(doc, "7. Reporting & Analytics")
add_text_area(doc, "Describe required dashboards: enrolment, completion, payouts, engagement by org, cooperative performance, etc.")

add_section_heading(doc, "8. User Stories & Acceptance Criteria")
add_text_area(doc, "Include user stories per role with clear acceptance tests, e.g. 'As an Org Admin...'. Link to JIRA or backlog IDs if applicable.")

add_section_heading(doc, "9. Risks & Assumptions")
add_text_area(doc, "List implementation risks (connectivity, adoption, compliance) and assumptions validated from Discovery.")

add_section_heading(doc, "10. Next Steps & Sign-off")
add_text_area(doc, "Outline review steps, approvals required, and transition into Phase 3: System Design & Prototyping.")

# Ensure output directory exists
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)

# Save
doc.save(OUTPUT_PATH)
print(f"✅ BRD Template created at:\n{OUTPUT_PATH}")

