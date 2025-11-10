import os
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- CONFIG ----------
ORG_NAME = "UNIQUE ENTREPRENEUR SOLUTIONS LIMITED"
FOOTER_TEXT = ("Omole Victoria Oluwatosin — Business Analyst | Data Analyst | "
               "Data Management Professional | Frontend & Backend Engineer")

BLUE_GREY = RGBColor(30, 55, 90)
LIGHT_GREY = "F0F0F0"

BASE_DIR = r"C:\Users\victo\kickoff_project\02_Phase2_Requirements_and_Governance"
os.makedirs(BASE_DIR, exist_ok=True)

WORD_PATH = os.path.join(BASE_DIR, "UniqueEntrepreneur_UserStories_and_Epics.docx")
EXCEL_PATH = os.path.join(BASE_DIR, "UniqueEntrepreneur_UserStories_and_Epics.xlsx")

# ---------- DATA ----------

epics = [
    ("E-01", "Multi-Tenant Org & School Management", "Setup and management of organisations, schools, and classes."),
    ("E-02", "Course Creation & Delivery", "Instructors create, manage, and publish multimedia courses."),
    ("E-03", "Learning & Assessment", "Students enrol, learn, take quizzes/assignments, and receive certificates."),
    ("E-04", "Analytics & Reporting", "Dashboards for engagement, completion, and financial insights."),
    ("E-05", "Governance & Compliance", "Implements RLS, GDPR/NDPR controls, and audit trails."),
    ("E-06", "Payments & Monetisation", "Course sales, coupons, payouts, and subscriptions."),
]

stories = [
    ("US-01", "Platform Admin",
     "create organisations so each has its own environment",
     "separate data and branding per organisation",
     "Org created with unique org_id, branding options, and default admin.", "High"),
    ("US-02", "Org Admin",
     "create schools and assign school admins",
     "delegate management within my organisation",
     "Schools linked to org_id; admins receive invitation email.", "High"),
    ("US-03", "Org Admin",
     "bulk-import teachers and students via CSV",
     "onboard faster without manual entry",
     "CSV validated; errors reported; valid rows created as users.", "Medium"),
    ("US-05", "Instructor",
     "upload videos, PDFs, and quizzes",
     "build engaging online courses",
     "Supports common formats; file size limits enforced; items previewable.", "High"),
    ("US-08", "Student",
     "enrol in a course and resume where I left off",
     "continue learning smoothly",
     "System stores last completed lesson and video position per course.", "High"),
    ("US-09", "Student",
     "take quizzes and see my results",
     "understand my performance",
     "Auto-graded quizzes; scores visible in course progress.", "High"),
    ("US-11", "Org Admin",
     "view a dashboard of learner activity",
     "monitor adoption and completion",
     "Shows active users, enrolments, completion %, top courses.", "High"),
    ("US-14", "DMP",
     "view audit logs of key actions",
     "ensure compliance and traceability",
     "Logs include who, what, when, before/after where relevant.", "High"),
    ("US-17", "Instructor",
     "set course prices and coupons",
     "monetise my content",
     "Price and discounts applied; integrated with payment gateway.", "High"),
]

# ---------- HELPER (WORD) ----------

def shade_cell(cell, fill_hex=LIGHT_GREY):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

# ---------- WORD DOCUMENT GENERATION ----------

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
style.font.size = Pt(10)

section = doc.sections[0]

# Header
header = section.header
if not header.paragraphs:
    header.add_paragraph()
hp = header.paragraphs[0]
hp.text = ""
hr = hp.add_run(ORG_NAME)
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
hr.font.size = Pt(11)
hr.font.bold = True
hr.font.color.rgb = BLUE_GREY

# Footer
footer = section.footer
if not footer.paragraphs:
    footer.add_paragraph()
fp = footer.paragraphs[0]
fp.text = ""
fr = fp.add_run(FOOTER_TEXT)
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr.font.size = Pt(9)
fr.font.color.rgb = BLUE_GREY

# Title
p = doc.add_paragraph()
r = p.add_run("Phase 2.2 — User Stories & Epics")
r.font.size = Pt(16)
r.font.bold = True
r.font.color.rgb = BLUE_GREY
doc.add_paragraph("This appendix defines epics and user stories for the Unique Entrepreneur Literacy Hub MVP.").alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_paragraph()

# Epics table
title_para = doc.add_paragraph()
tr = title_para.add_run("Epics Overview")
tr.font.bold = True
tr.font.size = Pt(12)
tr.font.color.rgb = BLUE_GREY

epic_table = doc.add_table(rows=1, cols=3)
headers = ["Epic ID", "Epic Title", "Description"]
for idx, h in enumerate(headers):
    cell = epic_table.rows[0].cells[idx]
    run = cell.paragraphs[0].add_run(h)
    run.font.bold = True
    run.font.color.rgb = BLUE_GREY

for epic in epics:
    row_cells = epic_table.add_row().cells
    for i, v in enumerate(epic):
        para = row_cells[i].paragraphs[0]
        para.add_run(v)
        shade_cell(row_cells[i])

doc.add_paragraph()

# User Stories table
us_title = doc.add_paragraph()
usr = us_title.add_run("User Stories")
usr.font.bold = True
usr.font.size = Pt(12)
usr.font.color.rgb = BLUE_GREY

us_table = doc.add_table(rows=1, cols=6)
us_headers = ["Story ID", "Role", "Action", "Goal", "Acceptance Criteria", "Priority"]
for idx, h in enumerate(us_headers):
    cell = us_table.rows[0].cells[idx]
    run = cell.paragraphs[0].add_run(h)
    run.font.bold = True
    run.font.color.rgb = BLUE_GREY

for s in stories:
    row = us_table.add_row().cells
    for i, v in enumerate(s):
        para = row[i].paragraphs[0]
        para.add_run(v)
        shade_cell(row[i])

doc.save(WORD_PATH)
print(f"✅ Word document created at:\n{WORD_PATH}")

# ---------- EXCEL GENERATION ----------

wb = Workbook()
ws = wb.active
ws.title = "UserStories"

excel_headers = [
    "Epic",
    "Story ID",
    "Role",
    "Action",
    "Goal",
    "Acceptance Criteria",
    "Priority",
    "Story Points",
    "Sprint",
    "Status"
]

header_font = Font(bold=True, color="FFFFFF")
header_fill = PatternFill(start_color="1E375A", end_color="1E375A", fill_type="solid")
header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)

# Header row
for col, h in enumerate(excel_headers, start=1):
    c = ws.cell(row=1, column=col, value=h)
    c.font = header_font
    c.fill = header_fill
    c.alignment = header_align

# Map stories to epics (simple round-robin just for template)
for i, s in enumerate(stories, start=2):
    epic_id = epics[(i - 2) % len(epics)][0]  # cycle epics
    row_values = [
        epic_id,
        s[0],
        s[1],
        s[2],
        s[3],
        s[4],
        s[5],
        "",   # Story Points (fillable)
        "",   # Sprint (fillable)
        ""    # Status (fillable)
    ]
    for col, v in enumerate(row_values, start=1):
        cell = ws.cell(row=i, column=col, value=v)
        cell.alignment = Alignment(wrap_text=True, vertical="top")

# Auto-width
for col in ws.columns:
    max_len = 0
    col_letter = col[0].column_letter
    for cell in col:
        if cell.value:
            max_len = max(max_len, len(str(cell.value)))
    ws.column_dimensions[col_letter].width = min(max_len + 2, 35)

wb.save(EXCEL_PATH)
print(f"✅ Excel workbook created at:\n{EXCEL_PATH}")
